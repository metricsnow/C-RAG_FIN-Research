"""
Conversation export utility module.

Provides functions to export conversation history to various formats:
JSON, Markdown, plain text, PDF, Word (DOCX), and CSV.
"""

import csv
import io
import json
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional
from uuid import uuid4

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from app.utils.logger import get_logger

logger = get_logger(__name__)


def export_to_json(
    messages: List[Dict[str, Any]],
    model: Optional[str] = None,
    conversation_id: Optional[str] = None,
) -> str:
    """
    Export conversation history to JSON format.

    Args:
        messages: List of message dictionaries with 'role', 'content',
            and optionally 'sources'
        model: Model name used for the conversation (optional)
        conversation_id: Unique conversation identifier (optional,
            generated if not provided)

    Returns:
        JSON string representation of the conversation
    """
    if conversation_id is None:
        conversation_id = str(uuid4())

    # Format messages for export
    export_messages = []
    for msg in messages:
        export_msg = {
            "role": msg.get("role", "unknown"),
            "content": msg.get("content", ""),
        }

        # Add timestamp if available
        if "timestamp" in msg:
            export_msg["timestamp"] = msg["timestamp"]

        # Add sources for assistant messages
        if msg.get("role") == "assistant" and "sources" in msg:
            export_msg["sources"] = msg["sources"]

        export_messages.append(export_msg)

    # Create export structure
    export_data = {
        "conversation_id": conversation_id,
        "created_at": datetime.now().isoformat(),
        "messages": export_messages,
        "metadata": {
            "model": model or "unknown",
            "total_messages": len(messages),
            "user_messages": sum(1 for m in messages if m.get("role") == "user"),
            "assistant_messages": sum(
                1 for m in messages if m.get("role") == "assistant"
            ),
        },
    }

    return json.dumps(export_data, indent=2, ensure_ascii=False)


def export_to_markdown(
    messages: List[Dict[str, Any]],
    model: Optional[str] = None,
    conversation_id: Optional[str] = None,
) -> str:
    """
    Export conversation history to Markdown format.

    Args:
        messages: List of message dictionaries with 'role', 'content',
            and optionally 'sources'
        model: Model name used for the conversation (optional)
        conversation_id: Unique conversation identifier (optional)

    Returns:
        Markdown string representation of the conversation
    """
    if conversation_id is None:
        conversation_id = str(uuid4())

    # Header
    lines = ["# Conversation Export"]
    lines.append(f"**Date:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"**Model:** {model or 'unknown'}")
    lines.append(f"**Conversation ID:** {conversation_id}")
    lines.append("")
    lines.append("## Messages")
    lines.append("")

    # Format messages
    for msg in messages:
        role = msg.get("role", "unknown")
        content = msg.get("content", "")

        # Add timestamp if available
        timestamp_str = ""
        if "timestamp" in msg:
            timestamp_str = f" *({msg['timestamp']})*"

        # Format based on role
        if role == "user":
            lines.append(f"### User{timestamp_str}")
            lines.append("")
            lines.append(content)
        elif role == "assistant":
            lines.append(f"### Assistant{timestamp_str}")
            lines.append("")
            lines.append(content)

            # Add sources if available
            if "sources" in msg and msg["sources"]:
                sources = msg["sources"]
                source_names = []
                for source in sources:
                    filename = source.get("filename") or source.get("source", "unknown")
                    # Extract just the filename if it's a path
                    if isinstance(filename, str) and "/" in filename:
                        filename = Path(filename).name
                    source_names.append(filename)

                if source_names:
                    lines.append("")
                    if len(source_names) == 1:
                        lines.append(f"**Source:** {source_names[0]}")
                    else:
                        lines.append(f"**Sources:** {', '.join(source_names)}")

        lines.append("")
        lines.append("---")
        lines.append("")

    return "\n".join(lines)


def export_to_txt(
    messages: List[Dict[str, Any]],
    model: Optional[str] = None,
    conversation_id: Optional[str] = None,
) -> str:
    """
    Export conversation history to plain text format.

    Args:
        messages: List of message dictionaries with 'role', 'content',
            and optionally 'sources'
        model: Model name used for the conversation (optional)
        conversation_id: Unique conversation identifier (optional)

    Returns:
        Plain text string representation of the conversation
    """
    if conversation_id is None:
        conversation_id = str(uuid4())

    # Header
    lines = ["Conversation Export"]
    lines.append("=" * 50)
    lines.append(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"Model: {model or 'unknown'}")
    lines.append(f"Conversation ID: {conversation_id}")
    lines.append("=" * 50)
    lines.append("")

    # Format messages
    for idx, msg in enumerate(messages, 1):
        role = msg.get("role", "unknown")
        content = msg.get("content", "")

        # Add timestamp if available
        timestamp_str = ""
        if "timestamp" in msg:
            timestamp_str = f" ({msg['timestamp']})"

        # Format based on role
        lines.append(f"[{idx}] {role.upper()}{timestamp_str}")
        lines.append("-" * 50)
        lines.append(content)

        # Add sources if available
        if role == "assistant" and "sources" in msg and msg["sources"]:
            sources = msg["sources"]
            source_names = []
            for source in sources:
                filename = source.get("filename") or source.get("source", "unknown")
                # Extract just the filename if it's a path
                if isinstance(filename, str) and "/" in filename:
                    filename = Path(filename).name
                source_names.append(filename)

            if source_names:
                lines.append("")
                if len(source_names) == 1:
                    lines.append(f"Source: {source_names[0]}")
                else:
                    lines.append(f"Sources: {', '.join(source_names)}")

        lines.append("")
        lines.append("")

    return "\n".join(lines)


def export_to_pdf(
    messages: List[Dict[str, Any]],
    model: Optional[str] = None,
    conversation_id: Optional[str] = None,
) -> bytes:
    """
    Export conversation history to PDF format.

    Args:
        messages: List of message dictionaries with 'role', 'content',
            and optionally 'sources'
        model: Model name used for the conversation (optional)
        conversation_id: Unique conversation identifier (optional)

    Returns:
        PDF file as bytes

    Raises:
        ImportError: If reportlab is not installed
    """
    if not REPORTLAB_AVAILABLE:
        raise ImportError(
            "reportlab is required for PDF export. "
            "Install it with: pip install reportlab"
        )

    if conversation_id is None:
        conversation_id = str(uuid4())

    # Create PDF in memory
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    story = []
    styles = getSampleStyleSheet()

    # Header
    header_style = styles["Heading1"]
    story.append(Paragraph("Conversation Export", header_style))
    story.append(Spacer(1, 0.2 * inch))

    # Metadata
    normal_style = styles["Normal"]
    story.append(
        Paragraph(
            f"<b>Date:</b> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            normal_style,
        )
    )
    story.append(Paragraph(f"<b>Model:</b> {model or 'unknown'}", normal_style))
    story.append(
        Paragraph(f"<b>Conversation ID:</b> {conversation_id[:8]}", normal_style)
    )
    story.append(Spacer(1, 0.3 * inch))

    # Messages
    heading_style = styles["Heading2"]
    for idx, msg in enumerate(messages, 1):
        role = msg.get("role", "unknown")
        content = msg.get("content", "")

        # Add timestamp if available
        timestamp_str = ""
        if "timestamp" in msg:
            timestamp_str = f" ({msg['timestamp']})"

        # Format role header
        story.append(Paragraph(f"[{idx}] {role.upper()}{timestamp_str}", heading_style))
        story.append(Spacer(1, 0.1 * inch))

        # Add content (escape HTML and handle line breaks)
        content_escaped = (
            content.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        )
        content_escaped = content_escaped.replace("\n", "<br/>")
        story.append(Paragraph(content_escaped, normal_style))

        # Add sources if available
        if role == "assistant" and "sources" in msg and msg["sources"]:
            sources = msg["sources"]
            source_names = []
            for source in sources:
                filename = source.get("filename") or source.get("source", "unknown")
                if isinstance(filename, str) and "/" in filename:
                    filename = Path(filename).name
                source_names.append(filename)

            if source_names:
                story.append(Spacer(1, 0.05 * inch))
                sources_text = (
                    f"<b>Source:</b> {source_names[0]}"
                    if len(source_names) == 1
                    else f"<b>Sources:</b> {', '.join(source_names)}"
                )
                story.append(Paragraph(sources_text, normal_style))

        story.append(Spacer(1, 0.2 * inch))

    # Build PDF
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def export_to_docx(
    messages: List[Dict[str, Any]],
    model: Optional[str] = None,
    conversation_id: Optional[str] = None,
) -> bytes:
    """
    Export conversation history to Word (DOCX) format.

    Args:
        messages: List of message dictionaries with 'role', 'content',
            and optionally 'sources'
        model: Model name used for the conversation (optional)
        conversation_id: Unique conversation identifier (optional)

    Returns:
        DOCX file as bytes

    Raises:
        ImportError: If python-docx is not installed
    """
    if not DOCX_AVAILABLE:
        raise ImportError(
            "python-docx is required for Word export. "
            "Install it with: pip install python-docx"
        )

    if conversation_id is None:
        conversation_id = str(uuid4())

    # Create document
    doc = Document()

    # Header
    doc.add_heading("Conversation Export", 0)
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Model: {model or 'unknown'}")
    doc.add_paragraph(f"Conversation ID: {conversation_id[:8]}")
    doc.add_paragraph("")  # Empty line

    # Messages
    for idx, msg in enumerate(messages, 1):
        role = msg.get("role", "unknown")
        content = msg.get("content", "")

        # Add timestamp if available
        timestamp_str = ""
        if "timestamp" in msg:
            timestamp_str = f" ({msg['timestamp']})"

        # Add role heading
        doc.add_heading(f"[{idx}] {role.upper()}{timestamp_str}", level=2)

        # Add content
        doc.add_paragraph(content)

        # Add sources if available
        if role == "assistant" and "sources" in msg and msg["sources"]:
            sources = msg["sources"]
            source_names = []
            for source in sources:
                filename = source.get("filename") or source.get("source", "unknown")
                if isinstance(filename, str) and "/" in filename:
                    filename = Path(filename).name
                source_names.append(filename)

            if source_names:
                source_text = (
                    f"Source: {source_names[0]}"
                    if len(source_names) == 1
                    else f"Sources: {', '.join(source_names)}"
                )
                p = doc.add_paragraph()
                run = p.add_run(source_text)
                run.bold = True

        doc.add_paragraph("")  # Empty line between messages

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_to_csv(
    messages: List[Dict[str, Any]],
    model: Optional[str] = None,
    conversation_id: Optional[str] = None,
) -> str:
    """
    Export conversation history to CSV format.

    Args:
        messages: List of message dictionaries with 'role', 'content',
            and optionally 'sources'
        model: Model name used for the conversation (optional)
        conversation_id: Unique conversation identifier (optional)

    Returns:
        CSV string representation of the conversation
    """
    if conversation_id is None:
        conversation_id = str(uuid4())

    # Create CSV in memory
    output = io.StringIO()
    writer = csv.writer(output)

    # Header row
    writer.writerow(
        [
            "Message Number",
            "Role",
            "Content",
            "Sources",
            "Timestamp",
            "Model",
            "Conversation ID",
        ]
    )

    # Data rows
    for idx, msg in enumerate(messages, 1):
        role = msg.get("role", "unknown")
        content = msg.get("content", "")

        # Extract sources
        sources = ""
        if role == "assistant" and "sources" in msg and msg["sources"]:
            source_names = []
            for source in msg["sources"]:
                filename = source.get("filename") or source.get("source", "unknown")
                if isinstance(filename, str) and "/" in filename:
                    filename = Path(filename).name
                source_names.append(filename)
            sources = ", ".join(source_names)

        timestamp = msg.get("timestamp", "")

        writer.writerow(
            [
                idx,
                role,
                content,
                sources,
                timestamp,
                model or "unknown",
                conversation_id[:8],
            ]
        )

    return output.getvalue()


def generate_export_filename(
    format_type: str, conversation_id: Optional[str] = None
) -> str:
    """
    Generate a filename for the exported conversation.

    Args:
        format_type: Export format ('json', 'markdown', 'txt')
        conversation_id: Unique conversation identifier (optional)

    Returns:
        Filename string with timestamp and format extension
    """
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    conv_id_short = (
        conversation_id[:8] if conversation_id else "conversation"
    )  # Use first 8 chars of UUID

    format_extensions = {
        "json": "json",
        "markdown": "md",
        "txt": "txt",
        "pdf": "pdf",
        "docx": "docx",
        "word": "docx",
        "csv": "csv",
    }

    extension = format_extensions.get(format_type.lower(), "txt")
    return f"conversation_{conv_id_short}_{timestamp}.{extension}"


def export_conversation(
    messages: List[Dict[str, Any]],
    format_type: str,
    model: Optional[str] = None,
    conversation_id: Optional[str] = None,
) -> tuple[str | bytes, str]:
    """
    Export conversation history to specified format.

    Args:
        messages: List of message dictionaries
        format_type: Export format ('json', 'markdown', 'txt', 'pdf', 'docx', 'csv')
        model: Model name used for the conversation (optional)
        conversation_id: Unique conversation identifier (optional)

    Returns:
        Tuple of (exported_content, filename)
        - For text formats (json, markdown, txt, csv): content is str
        - For binary formats (pdf, docx): content is bytes

    Raises:
        ValueError: If format is unsupported or messages are empty
        ImportError: If required library is not installed (for PDF/DOCX)
    """
    if not messages:
        raise ValueError("Cannot export empty conversation")

    format_type_lower = format_type.lower()

    if format_type_lower == "json":
        content = export_to_json(messages, model, conversation_id)
    elif format_type_lower in ["markdown", "md"]:
        content = export_to_markdown(messages, model, conversation_id)
    elif format_type_lower in ["txt", "text"]:
        content = export_to_txt(messages, model, conversation_id)
    elif format_type_lower == "pdf":
        content = export_to_pdf(messages, model, conversation_id)
    elif format_type_lower in ["docx", "word"]:
        content = export_to_docx(messages, model, conversation_id)
    elif format_type_lower == "csv":
        content = export_to_csv(messages, model, conversation_id)
    else:
        raise ValueError(
            f"Unsupported format type: {format_type}. "
            f"Supported formats: json, markdown, txt, pdf, docx, csv"
        )

    filename = generate_export_filename(format_type_lower, conversation_id)

    return content, filename


def batch_export_conversations(
    conversations: List[Dict[str, Any]],
    format_type: str,
    model: Optional[str] = None,
) -> List[tuple[str | bytes, str]]:
    """
    Export multiple conversations to specified format.

    Args:
        conversations: List of conversation dictionaries, each containing:
            - 'messages': List of message dictionaries
            - 'conversation_id': Optional conversation identifier
            - 'model': Optional model name (overrides global model parameter)
        format_type: Export format ('json', 'markdown', 'txt', 'pdf', 'docx', 'csv')
        model: Default model name for conversations without model specified

    Returns:
        List of tuples (exported_content, filename) for each conversation

    Raises:
        ValueError: If format is unsupported or conversations list is empty
    """
    if not conversations:
        raise ValueError("Cannot export empty conversations list")

    results = []
    for conv in conversations:
        conv_messages = conv.get("messages", [])
        if not conv_messages:
            logger.warning("Skipping conversation with no messages")
            continue

        conv_id = conv.get("conversation_id")
        conv_model = conv.get("model", model)

        try:
            content, filename = export_conversation(
                conv_messages, format_type, conv_model, conv_id
            )
            results.append((content, filename))
        except Exception as e:
            logger.error(f"Failed to export conversation {conv_id}: {e}")
            continue

    return results
